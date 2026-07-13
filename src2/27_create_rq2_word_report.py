"""Create a Word report for RQ2 methodology and results."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "rq2" / "documents" / "RQ2_Methodology_and_Results.docx"

SUMMARY_TASK = ROOT / "rq2" / "outputs" / "rq2_stats_descriptive_by_task.csv"
SUMMARY_PERT = ROOT / "rq2" / "outputs" / "rq2_stats_descriptive_by_perturbation.csv"
SUMMARY_DRIFT_PERF = ROOT / "rq2" / "outputs" / "rq2_formal_available_drift_performance_summary.csv"
PERF_SUMMARY = ROOT / "rq2" / "outputs" / "rq2_formal_available_performance_change_summary.csv"
CORR = ROOT / "rq2" / "outputs" / "rq2_stats_correlations.csv"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
BORDER = "B8C2CC"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as file:
        return list(csv.DictReader(file))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def format_table(table, widths: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    if widths:
        set_table_width(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, GRAY_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table, widths)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level in {1, 2} else DARK_BLUE
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_source_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def pct(value: str) -> str:
    if value == "":
        return ""
    return f"{float(value):.6f}"


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("RQ2 Methodology and Results")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Semantic Drift and Task Performance Change Under Prompt Perturbation")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(89, 89, 89)

    add_heading(doc, "1. Research Question", 1)
    add_body(
        doc,
        "RQ2 examines whether output-level semantic drift is associated with task performance change in tasks with objective evaluation criteria. Operationally, the analysis asks whether greater semantic drift after prompt perturbation corresponds to a larger performance decrease in factual question answering, mathematical reasoning, and code generation, and whether this relationship differs across task types.",
    )
    add_body(
        doc,
        "Open-ended writing is excluded from RQ2 because it lacks a stable objective correctness criterion. Its outputs remain relevant for RQ1 semantic drift analysis, but they are not included in the performance-change analysis.",
    )

    add_heading(doc, "2. Data and Analysis Unit", 1)
    add_body(
        doc,
        "The RQ2 analysis uses the formal RQ1 original and perturbed generation files. The RQ2 subset contains 150 original outputs and 750 perturbed outputs: three task types, ten items per task type, five perturbation types, and five repeated generations per prompt condition.",
    )
    add_body(
        doc,
        "The main analysis unit is one item-perturbation pair. This creates 150 item-level observations: 3 task types x 10 items x 5 perturbation types. For each observation, the analysis combines an original performance estimate, a perturbed performance estimate, an absolute performance-change score, a PDR value, and a noise-corrected semantic drift value.",
    )

    add_heading(doc, "3. Task-Specific Performance Metrics", 1)
    add_body(
        doc,
        "RQ2 does not force all tasks into a single binary correctness metric. Instead, performance is defined according to the evaluation logic of each task type.",
    )
    add_table(
        doc,
        ["Task type", "Performance metric", "Interpretation"],
        [
            ["factual_qa", "Continuous factual QA score", "Normalized reference-answer containment first; SQuAD-style token F1 backup. No binary correctness label is assigned."],
            ["math_reasoning", "Final-answer correctness, 1/0", "The extracted final answer is compared with the reference answer using normalized and symbolic equivalence."],
            ["code_generation", "Unit-test pass/fail, 1/0", "Generated Python code is executed against HumanEvalPack tests; passing all tests receives score 1."],
        ],
        [1.35, 2.15, 3.0],
    )
    add_source_note(doc, "Table 1. Task-specific performance definitions used in RQ2.")

    add_heading(doc, "3.1 Factual QA", 2)
    add_body(
        doc,
        "Factual QA uses a SQuAD-style automatic lexical score adapted for full-sentence LLM outputs. The evaluator first normalizes the reference answer and model output by lowercasing, removing punctuation, removing English articles, and collapsing whitespace. If the normalized reference answer appears in the normalized output, the factual performance score is set to 1.0. Otherwise, the evaluator computes token-level F1 between the normalized reference answer and the full output.",
    )
    add_body(
        doc,
        "This design avoids penalizing long but correct full-sentence answers when the reference answer is explicitly present. It also avoids imposing an arbitrary binary correctness threshold on factual QA. The limitation is that the score remains lexical and may underestimate semantically correct answers that do not share tokens with the reference answer.",
    )

    add_heading(doc, "3.2 Mathematical Reasoning", 2)
    add_body(
        doc,
        "Mathematical reasoning is evaluated through final-answer correctness. The evaluator extracts a final answer from LaTeX boxed answers, final-answer phrases, or the last simple numeric expression. It then compares the extracted answer with the reference using normalized string matching and symbolic or numeric equivalence when possible. Outputs that cannot be extracted or compared automatically are treated as incorrect under the fully automatic conservative policy.",
    )

    add_heading(doc, "3.3 Code Generation", 2)
    add_body(
        doc,
        "Code generation is evaluated by functional correctness. The evaluator extracts Python code from each generated output, loads the corresponding HumanEvalPack test metadata, and executes the candidate solution in a temporary subprocess. Passing all tests receives a score of 1; syntax errors, runtime errors, assertion failures, and timeouts receive a score of 0.",
    )

    add_heading(doc, "4. Semantic Drift and Performance Change", 1)
    add_body(
        doc,
        "The semantic drift variable is noise-corrected using the RQ1 repeated-sampling baseline. For each item, baseline similarity is the average pairwise Sentence-BERT cosine similarity among the five original-prompt outputs. Perturbation similarity is the average similarity between the five original-prompt outputs and the five perturbed-prompt outputs. Noise-corrected drift is then calculated as baseline similarity minus perturbation similarity.",
    )
    add_body(
        doc,
        "Performance change is calculated as original performance minus perturbed performance. A positive value means that performance decreased after perturbation. The study also reports PDR, defined as performance change divided by original performance when original performance is greater than zero.",
    )

    add_heading(doc, "5. Formal Results", 1)
    add_heading(doc, "5.1 Original and Perturbed Performance by Task", 2)
    task_rows = read_csv(SUMMARY_TASK)
    task_table = []
    for row in task_rows:
        task_table.append(
            [
                row["task_type"],
                pct(row["absolute_performance_change_mean"]),
                f"[{float(row['absolute_performance_change_ci95_low']):.6f}, {float(row['absolute_performance_change_ci95_high']):.6f}]",
                pct(row["noise_corrected_drift_mean"]),
                pct(row["pdr_mean"]),
            ]
        )
    add_table(
        doc,
        ["Task type", "Mean performance change", "95% CI", "Mean drift", "Mean PDR"],
        task_table,
        [1.45, 1.35, 1.7, 1.0, 1.0],
    )
    add_source_note(doc, "Table 2. RQ2 results summarized by task type. Performance change is original performance minus perturbed performance.")

    add_body(
        doc,
        "All three objective task types show lower average performance after perturbation. Factual QA shows the largest mean performance decrease, followed closely by code generation. Math reasoning shows a smaller mean decrease, and its confidence interval includes zero.",
    )

    add_heading(doc, "5.2 Results by Perturbation Type", 2)
    pert_rows = read_csv(SUMMARY_PERT)
    order = {"paraphrasing": 0, "reordering": 1, "formatting_changes": 2, "surface_noise": 3, "context_injection": 4}
    pert_table = []
    for row in sorted(pert_rows, key=lambda r: order.get(r["perturbation_type"], 99)):
        pert_table.append(
            [
                row["perturbation_type"],
                pct(row["noise_corrected_drift_mean"]),
                f"[{float(row['noise_corrected_drift_ci95_low']):.6f}, {float(row['noise_corrected_drift_ci95_high']):.6f}]",
                pct(row["absolute_performance_change_mean"]),
                f"[{float(row['absolute_performance_change_ci95_low']):.6f}, {float(row['absolute_performance_change_ci95_high']):.6f}]",
            ]
        )
    add_table(
        doc,
        ["Perturbation type", "Mean drift", "95% CI drift", "Mean performance change", "95% CI change"],
        pert_table,
        [1.55, 1.0, 1.45, 1.35, 1.45],
    )
    add_source_note(doc, "Table 3. RQ2 results summarized by perturbation type.")
    add_body(
        doc,
        "Paraphrasing produces the largest average semantic drift and the largest average performance decrease. Reordering and formatting changes show moderate effects, while context injection and surface noise show smaller average performance effects in this run.",
    )

    add_heading(doc, "5.3 Task-by-Perturbation Results", 2)
    drift_rows = read_csv(SUMMARY_DRIFT_PERF)
    drift_table = []
    for row in drift_rows:
        drift_table.append(
            [
                row["task_type"],
                row["perturbation_type"],
                pct(row["mean_noise_corrected_drift"]),
                pct(row["mean_absolute_performance_change"]),
                pct(row["pearson_drift_performance_change"]),
            ]
        )
    add_table(
        doc,
        ["Task type", "Perturbation", "Mean drift", "Mean perf. change", "Pearson r"],
        drift_table,
        [1.25, 1.45, 1.0, 1.25, 1.0],
    )
    add_source_note(doc, "Table 4. Task-by-perturbation drift-performance results. Each Pearson value is descriptive because each cell contains ten items.")

    add_heading(doc, "6. Statistical Analysis", 1)
    add_body(
        doc,
        "The statistical analysis uses absolute performance change as the outcome and noise-corrected semantic drift as the main predictor. Task type and perturbation type are treated as categorical factors. Correlation tests were used for association summaries, and OLS regressions with HC3 robust standard errors were used for inferential models.",
    )

    add_heading(doc, "6.1 Correlation Results", 2)
    corr_rows = read_csv(CORR)
    corr_table = []
    for row in corr_rows:
        if row.get("group") == "overall" or (row.get("task_type") and not row.get("perturbation_type")):
            label = row.get("group") or row.get("task_type")
            corr_table.append(
                [
                    label,
                    row["n"],
                    pct(row["pearson_r"]),
                    pct(row["pearson_p"]),
                    pct(row["spearman_rho"]),
                    pct(row["spearman_p"]),
                ]
            )
    add_table(
        doc,
        ["Scope", "n", "Pearson r", "Pearson p", "Spearman rho", "Spearman p"],
        corr_table,
        [1.35, 0.55, 1.0, 1.0, 1.15, 1.0],
    )
    add_source_note(doc, "Table 5. Overall and task-level association between noise-corrected drift and performance change.")
    add_body(
        doc,
        "The pooled Pearson correlation is positive and statistically significant, while the pooled Spearman correlation is not significant. This suggests a linear association in the pooled data but not a stable universal monotonic pattern. At the task level, factual QA shows the most consistent relationship: both Pearson and Spearman correlations are positive and statistically significant. Code generation shows a significant Pearson correlation but not a significant Spearman correlation. Math reasoning shows a negative Pearson correlation, indicating a different relationship between embedding-based drift and final-answer correctness change.",
    )

    add_heading(doc, "6.2 Regression Results", 2)
    add_table(
        doc,
        ["Model", "Specification", "Key result", "Interpretation"],
        [
            [
                "OLS 1",
                "performance_change ~ drift",
                "drift coef. = 1.6937, p = 0.007, R2 = 0.107",
                "Higher drift is associated with larger performance decrease in pooled data.",
            ],
            [
                "OLS 2",
                "performance_change ~ drift * task_type",
                "math interaction = -8.4603, p = 0.002, R2 = 0.220",
                "The drift-performance relationship differs by task type, especially for math reasoning.",
            ],
            [
                "OLS 3",
                "performance_change ~ drift * task_type + perturbation_type",
                "paraphrasing coef. = 0.1491, p = 0.047; math interaction p = 0.004, R2 = 0.264",
                "Paraphrasing remains associated with larger performance decrease after controlling for task and drift.",
            ],
        ],
        [0.8, 1.75, 1.85, 2.1],
    )
    add_source_note(doc, "Table 6. Regression models with HC3 robust standard errors.")
    add_body(
        doc,
        "A mixed-effects robustness check with item-level random intercepts converged, but the random-effects covariance was singular and the estimated group variance was approximately zero. Therefore, it is treated as a robustness check rather than the main inferential result.",
    )

    add_heading(doc, "7. Answer to RQ2", 1)
    add_body(
        doc,
        "The current results suggest that semantic drift and task performance change are associated, but the association is task-dependent and perturbation-dependent rather than universal. The clearest positive pattern appears for paraphrasing, especially in factual QA and code generation. Math reasoning shows performance decreases even when embedding-based semantic drift is close to zero, suggesting that final-answer correctness can change without a large SBERT-measured semantic drift.",
    )
    add_bullet(doc, "Paraphrasing produced the largest average semantic drift and the largest average performance decrease.")
    add_bullet(doc, "Factual QA showed the most consistent drift-performance association in correlation analysis.")
    add_bullet(doc, "Code generation showed a strong linear association, especially under paraphrasing.")
    add_bullet(doc, "Math reasoning showed a different pattern: correctness can change even when semantic drift remains small.")

    add_heading(doc, "8. Limitations", 1)
    add_body(
        doc,
        "The sample contains 150 item-perturbation observations, with only ten items in each task-by-perturbation cell. Cell-level correlations should therefore be interpreted as descriptive rather than strong inferential evidence. The factual QA metric is automatic and lexical, so it may underestimate semantically correct answers that do not share tokens with the reference answer. The code metric depends on the available HumanEvalPack tests and does not prove correctness for all possible inputs. Finally, the semantic drift measure uses Sentence-BERT embeddings, which may not capture small final-answer changes in mathematical reasoning.",
    )

    add_heading(doc, "References", 1)
    references = [
        "K. Zhu et al., \"PromptBench: Towards evaluating the robustness of large language models on adversarial prompts,\" arXiv:2306.04528, 2023.",
        "A. Agrawal, L. Alazraki, S. Honarvar, and M. Rei, \"Enhancing LLM robustness to perturbed instructions: An empirical study,\" arXiv:2504.02733, 2025.",
        "A. Chatterjee et al., \"POSIX: A prompt sensitivity index for large language models,\" arXiv:2410.02185, 2024.",
        "F. Errica et al., \"What did I do wrong? Quantifying LLMs' sensitivity and consistency to prompt engineering,\" NAACL-HLT, 2025.",
        "J. Haase et al., \"Within-model vs. between-prompt variability in large language models for creative tasks,\" arXiv:2601.21339, 2026.",
        "N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence embeddings using Siamese BERT-networks,\" EMNLP-IJCNLP, 2019.",
        "P. Rajpurkar et al., \"SQuAD: 100,000+ questions for machine comprehension of text,\" EMNLP, 2016.",
    ]
    for ref in references:
        add_body(doc, ref)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("RQ2 Methodology and Results").font.size = Pt(9)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_document()
